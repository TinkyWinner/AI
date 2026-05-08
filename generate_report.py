#!/usr/bin/env python3
"""
Генератор отчёта в Word формате для работы по CNN классификации подписей.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Добавляет заголовок."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    """Добавляет параграф с опциональным форматированием."""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def create_report():
    """Создаёт Word документ с отчётом."""
    doc = Document()
    
    # Титульная страница
    title = doc.add_heading('Практическое задание №3', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Применение CNN для решения задачи классификации изображений')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    
    topic = doc.add_paragraph('Тема: Система обнаружения поддельных подписей на основе CNN')
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Пустая строка
    
    info = doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. Введение (Постановка задачи)
    add_heading(doc, '1. Введение (Постановка задачи)', level=1)
    
    doc.add_paragraph(
        'Целью работы является построение модели глубокого обучения, которая по изображению подписи '
        'определяет, является ли она подлинной или содержит признаки подделки, монтажа или фальсификации. '
        'Для таких задач важно не только распознать крупные формы, но и уловить мелкие текстурные артефакты, '
        'следы сжатия, неоднородности печати и аномалии в локальных областях изображения.'
    )
    
    doc.add_paragraph(
        'Задача является задачей бинарной классификации: каждому изображению подписи назначается один из двух '
        'классов — "подлинная" (real) или "поддельная" (fake). Модель должна быть достаточно чувствительна к '
        'качеству печати, артефактам сжатия и другим характерным признакам подделок, но вместе с тем устойчива '
        'к вариативности подлинных подписей (разные люди, условия съёмки, качество сканирования).'
    )
    
    doc.add_page_break()
    
    # 2. Обзор данных
    add_heading(doc, '2. Обзор данных', level=1)
    
    add_heading(doc, '2.1 Описание датасета', level=2)
    doc.add_paragraph(
        'В работе используется датасет подписей Signature Verification Dataset (Kaggle, '
        '"robinreni/signature-verification-dataset"). Датасет содержит изображения подлинных и '
        'поддельных/сфальсифицированных подписей, разделённые на обучающую (train) и тестовую (test) части.'
    )
    
    add_heading(doc, '2.2 Баланс классов', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выборка'
    hdr_cells[1].text = 'Real (подлинные)'
    hdr_cells[2].text = 'Fake (поддельные)'
    
    # Данные
    data = [
        ('train', '703 (53.3%)', '617 (46.7%)'),
        ('val', '184 (55.9%)', '145 (44.1%)'),
        ('test', '252 (50.4%)', '248 (49.6%)'),
    ]
    
    for i, (split_name, real, fake) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = split_name
        row_cells[1].text = real
        row_cells[2].text = fake
    
    doc.add_paragraph(
        '\nКлассы в датасете практически сбалансированы (≈50/50), что позволяет использовать accuracy '
        'в качестве основной метрики и не требует специальной обработки дисбаланса. Однако в контексте '
        'безопасности более важен recall для класса "fake" — пропуск подделки опаснее ложного срабатывания.'
    )
    
    add_heading(doc, '2.3 Предобработка и аугментация данных', level=2)
    doc.add_paragraph(
        'Аугментации применяются только к обучающей выборке для снижения переобучения и обучения модели '
        'инвариантности к трансформациям:'
    )
    
    aug_list = doc.add_paragraph(style='List Bullet')
    run_aug = aug_list.add_run('RandomResizedCrop')
    run_aug.bold = True
    aug_list.add_run(': случайное кадрирование и изменение масштаба (0.8–1.0)')
    
    aug_list = doc.add_paragraph(style='List Bullet')
    run_aug = aug_list.add_run('RandomHorizontalFlip')
    run_aug.bold = True
    aug_list.add_run(': зеркальное отражение с вероятностью 50%')
    
    aug_list = doc.add_paragraph(style='List Bullet')
    run_aug = aug_list.add_run('RandomRotation')
    run_aug.bold = True
    aug_list.add_run(': случайный поворот на ±10°')
    
    aug_list = doc.add_paragraph(style='List Bullet')
    run_aug = aug_list.add_run('ColorJitter')
    run_aug.bold = True
    aug_list.add_run(': варьирование яркости, контраста, насыщенности и оттенка')
    
    doc.add_paragraph(
        '\nНа валидационной и тестовой выборке применяется только нормализация без аугментаций, '
        'чтобы обеспечить честную оценку качества модели.'
    )
    
    doc.add_page_break()
    
    # 3. Методология
    add_heading(doc, '3. Методология', level=1)
    
    add_heading(doc, '3.1 Выбор архитектуры CNN и её обоснование', level=2)
    doc.add_paragraph(
        'В качестве базовой модели выбрана ResNet50 (Residual Networks), предобученная на датасете ImageNet. '
        'Причины выбора:'
    )
    
    reasons = [
        ('Остаточные связи (skip connections)', 
         'позволяют глубокой сети (50 слоёв) эффективно обучаться без деградации градиента, '
         'сохраняя информацию на различных уровнях абстракции'),
        ('Предобучение на ImageNet',
         'даёт полезные низкоуровневые признаки (края, линии, текстуры) и среднеуровневые паттерны, '
         'которые хорошо переносятся на задачу анализа подписей'),
        ('Выразительность',
         'позволяет извлекать и комбинировать мелкие различия между подлинными и поддельными изображениями'),
        ('Баланс',
         'обеспечивает хороший компромисс между точностью и вычислительной стоимостью в режиме transfer learning'),
    ]
    
    for title, desc in reasons:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    add_heading(doc, '3.2 Transfer Learning и Fine-tuning', level=2)
    doc.add_paragraph(
        'Обучение модели проходит в два этапа:'
    )
    
    add_heading(doc, 'Этап 1: Обучение классификатора (заморозка backbone)', level=3)
    doc.add_paragraph(
        'На первом этапе вес ResNet50 заморожены (requires_grad=False), обучаются только добавленные слои '
        '(классификатор). Это быстро адаптирует предобученные признаки к нашей задаче.', style='List Number'
    )
    
    doc.add_paragraph(
        'Архитектура классификатора:'
    )
    classifier_bullets = [
        'Linear(2048 → 256): линейное преобразование из признаков ResNet',
        'ReLU: активация для нелинейности',
        'Dropout(0.5): регуляризация для предотвращения переобучения',
        'Linear(256 → 1): выходной логит для бинарной классификации',
    ]
    for item in classifier_bullets:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading(doc, 'Этап 2: Fine-tuning (разморозка верхних слоёв)', level=3)
    doc.add_paragraph(
        'На втором этапе разморозивается верхний блок ResNet (layer4) вместе с классификатором. '
        'Learning rate устанавливается очень низким (1e-5) для тонкой доменной адаптации признаков '
        'без разрушения предобученных весов.'
    )
    
    add_heading(doc, '3.3 Функция потерь, оптимизатор и метрики', level=2)
    
    p = doc.add_paragraph(style='List Bullet')
    run_loss = p.add_run('BCEWithLogitsLoss: ')
    run_loss.bold = True
    p.add_run('численно стабильная функция потерь для бинарной классификации, объединяющая '
              'sigmoid активацию и binary cross-entropy')
    
    p = doc.add_paragraph(style='List Bullet')
    run_adam = p.add_run('AdamW: ')
    run_adam.bold = True
    p.add_run('оптимизатор с регуляризацией весов, подходит для глубоких сетей')
    
    doc.add_paragraph('Метрики для оценки:')
    metrics = [
        ('Accuracy', 'доля правильных предсказаний'),
        ('Precision', 'доля истинных подделок среди предсказанных подделок'),
        ('Recall', 'доля найденных подделок от всех реальных подделок; особенно важна для задачи безопасности'),
        ('F1-score', 'гармоническое среднее precision и recall'),
        ('Confusion Matrix', 'матрица ошибок для анализа типов ошибок (TP, FN, FP, TN)'),
    ]
    for metric, desc in metrics:
        p = doc.add_paragraph(style='List Bullet')
        run_metric = p.add_run(f'{metric}: ')
        run_metric.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 4. Эксперимент и результаты
    add_heading(doc, '4. Эксперимент и результаты', level=1)
    
    add_heading(doc, '4.1 Гиперпараметры обучения', level=2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[1].text = 'Значение'
    
    params = [
        ('Размер изображения', '224×224 px'),
        ('Batch size', '32'),
        ('Этап 1: эпохи', '5'),
        ('Этап 1: learning rate', '1e-3'),
        ('Этап 2: эпохи', '3'),
        ('Этап 2: learning rate', '1e-5'),
    ]
    
    for i, (param, value) in enumerate(params, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = param
        row_cells[1].text = value
    
    add_heading(doc, '4.2 Итоговые метрики на тестовой выборке', level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Метрика'
    hdr_cells[1].text = 'Значение'
    
    metrics_data = [
        ('Loss', '0.2121'),
        ('Accuracy', '0.914 (91.4%)'),
        ('Precision (fake)', '0.925 (92.5%)'),
        ('Recall (fake)', '0.899 (89.9%)'),
        ('F1-score (fake)', '0.912 (91.2%)'),
    ]
    
    for i, (metric, value) in enumerate(metrics_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    add_heading(doc, '4.3 Анализ матрицы ошибок', level=2)
    
    doc.add_paragraph(
        'На тестовой выборке из 500 примеров получены следующие результаты (при пороге принятия решения = 0.5):'
    )
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Предсказано: Real'
    hdr_cells[2].text = 'Предсказано: Fake'
    
    # Data
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Реально: Real'
    row1_cells[1].text = '234 (TN)'
    row1_cells[2].text = '18 (FP)'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Реально: Fake'
    row2_cells[1].text = '25 (FN)'
    row2_cells[2].text = '223 (TP)'
    
    doc.add_paragraph(
        '\nИнтерпретация:'
    )
    
    interpretation = [
        ('TP = 223', 'верно найдено подделок — основной успех модели'),
        ('FN = 25', 'пропущено подделок — представляет риск для системы безопасности'),
        ('FP = 18', 'ложные тревоги (подлинные классифицированы как поддельные) — раздражают пользователей, но менее опасны'),
        ('TN = 234', 'верно отклонено подлинных подписей'),
    ]
    
    for label, desc in interpretation:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(label).bold = True
        p.add_run(f': {desc}')
    
    doc.add_paragraph(
        '\nОсновное наблюдение: recall (89.9%) слегка ниже precision (92.5%), что означает, что модель '
        'чаще пропускает небольшую часть подделок (25 из 248), чем даёт ложные срабатывания. Для критичного '
        'приложения можно снизить порог решения (например, до 0.4–0.45) для увеличения recall, понимая, что '
        'произойдёт небольшой рост ложных тревог.'
    )
    
    add_heading(doc, '4.4 Качественный анализ ошибок', level=2)
    
    doc.add_paragraph(
        'Изучение ошибочных предсказаний показало следующие причины:'
    )
    
    errors_analysis = [
        ('Пропущенные подделки (FN)', 
         'часто имеют высокое качество печати, минимальные видимые артефакты на доступном разрешении, '
         'хороший контраст и четкие edges. Модель путает их с подлинными, так как их признаки близки к настоящим подписям.'),
        ('Ложные тревоги (FP)',
         'возникают при нестандартных условиях съёмки (плохое освещение, сильное сжатие, неоднородный фон), '
         'что может провоцировать артефакты, похожие на признаки подделки.'),
    ]
    
    for error_type, explanation in errors_analysis:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(error_type).bold = True
        p.add_run(f': {explanation}')
    
    doc.add_page_break()
    
    # 5. Выводы
    add_heading(doc, '5. Выводы и рекомендации', level=1)
    
    doc.add_paragraph(
        'В рамках работы успешно построена модель на базе ResNet50 для бинарной классификации подписей '
        '(подлинные vs поддельные) с использованием transfer learning и двухэтапного обучения.'
    )
    
    add_heading(doc, 'Ключевые результаты:', level=2)
    
    results = [
        'Достигнута точность 91.4% на тестовой выборке',
        'Recall для класса "fake" составляет 89.9%, что демонстрирует хорошую способность к выявлению подделок',
        'Precision 92.5% показывает, что большинство предсказанных подделок действительно являются поддельными',
        'Модель выучила релевантные признаки (видно из Grad-CAM): фокусируется на локальных области подписи, а не на фоне',
    ]
    
    for result in results:
        doc.add_paragraph(result, style='List Bullet')
    
    add_heading(doc, 'Почему ResNet50 и transfer learning эффективны:', level=2)
    
    reasons = [
        'Предобученные веса на ImageNet уже содержат общие визуальные признаки (текстуры, формы, края), '
        'которые хорошо переносятся на новую задачу',
        'Глубина сети (50 слоёв) с остаточными связями позволяет эффективно извлекать признаки на разных '
        'уровнях абстракции, необходимых для выявления мелких артефактов подделок',
        'Двухэтапное обучение (заморозка + fine-tuning) предотвращает переобучение на ограниченном датасете '
        'и позволяет адаптировать общие признаки под специфику задачи',
    ]
    
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    add_heading(doc, 'Практические рекомендации:', level=2)
    
    recommendations = [
        'Для приложений, где критичнее не пропустить подделку (банки, государственные учреждения), '
        'можно снизить порог решения до 0.4–0.45, чтобы повысить recall за счёт небольшого роста ложных тревог',
        'Для системы с низкой толерантностью к ложным срабатываниям (массовое использование, высокий трафик) '
        'можно оставить порог на уровне 0.5 и внедрить двухуровневую проверку (первый уровень — быстрая автоматическая '
        'классификация, второй — ручная проверка сомнительных случаев)',
        'Можно улучшить датасет, собрав больше примеров высококачественных подделок, чтобы модель лучше их распознавала',
        'Рассмотреть ансамбли моделей (например, ResNet50 + EfficientNet) для повешения надёжности',
    ]
    
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(rec, style='List Number')
    
    add_heading(doc, 'Заключение:', level=2)
    
    doc.add_paragraph(
        'Работа демонстрирует практическую применимость CNN и transfer learning для задач классификации '
        'изображений с ограниченным датасетом. Достигнутый уровень точности (91.4%) достаточен для использования '
        'в системах первичного скрининга подделок, с понимаем что критичные решения требуют дополнительной ручной проверки.'
    )
    
    # Сохранение документа
    output_path = '/workspaces/AI/ai/Отчёт_практическое_задание_CNN.docx'
    doc.save(output_path)
    print(f'Отчёт сохранён: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
