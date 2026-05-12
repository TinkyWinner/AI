#!/usr/bin/env python3
"""Generate a *full* NIR report (DOCX) with text, tables, and figures.

Inputs (existing files in this repo):
- R scripts: task1.R, task2_edited.R, task3.R (included in appendices)
- Notebooks: task4_variant17_colab.ipynb, task5_ccdata_pca.ipynb, task6_ccdata_clustering.ipynb
  (figures/tables extracted from embedded outputs)
- SPSS file: r13i_os26b.sav (used for Task 3 computations)

Because R is not available in this dev container, tasks 1–3 are reproduced in Python
using the same datasets/logic; the original R code is still attached in appendices.

This script follows the methodical guidelines from Методические_указания_по_выполнению_НИР_2024.docx:
- Title/assignment pages
- Table of contents placeholder
- Tasks 1–6 (each starts on a new page)
- Conclusion
- References (>=5)
- Appendices with code
"""

from __future__ import annotations

import argparse
import base64
import json
import re
import sys
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional


def ensure_package(import_name: str, pip_name: str | None = None) -> None:
    try:
        __import__(import_name)
    except ImportError:
        pkg = pip_name or import_name
        subprocess.check_call([sys.executable, "-m", "pip", "-q", "install", pkg])


# Runtime deps
ensure_package("numpy")
ensure_package("pandas")
ensure_package("docx", "python-docx")
ensure_package("statsmodels")
ensure_package("lxml")
ensure_package("pyreadstat")
ensure_package("matplotlib")
ensure_package("seaborn")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

import numpy as np  # noqa: E402
import pandas as pd  # noqa: E402

import matplotlib.pyplot as plt  # noqa: E402
import seaborn as sns  # noqa: E402

import statsmodels.api as sm  # noqa: E402
from statsmodels.stats.outliers_influence import variance_inflation_factor  # noqa: E402

import pyreadstat  # noqa: E402


@dataclass
class Counters:
    table_no: int = 0
    fig_no: int = 0


def safe_filename(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^0-9A-Za-zА-Яа-я_\-]+", "", s)
    return s


def set_base_styles(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 14) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)


def add_center_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc: Document, kind: str, no: int, text: str) -> None:
    p = doc.add_paragraph(f"{kind} {no} — {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, df: pd.DataFrame, title: str, counters: Counters, max_rows: int = 20) -> None:
    counters.table_no += 1
    add_caption(doc, "Таблица", counters.table_no, title)

    show = df.copy()
    if len(show) > max_rows:
        show = show.head(max_rows)

    show = show.reset_index(drop=False) if show.index.name or show.index.names != [None] else show

    n_rows = len(show) + 1
    n_cols = len(show.columns)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    for j, col in enumerate(show.columns):
        table.cell(0, j).text = str(col)

    for i in range(len(show)):
        for j, col in enumerate(show.columns):
            v = show.iloc[i, j]
            if isinstance(v, (float, np.floating)):
                text = f"{v:.4g}"
            else:
                text = str(v)
            table.cell(i + 1, j).text = text

    doc.add_paragraph("")


def add_picture(doc: Document, img_path: Path, title: str, counters: Counters, width_in: float = 6.2) -> None:
    counters.fig_no += 1
    doc.add_picture(str(img_path), width=Inches(width_in))
    add_caption(doc, "Рисунок", counters.fig_no, title)
    doc.add_paragraph("")


def add_code_block(doc: Document, code: str) -> None:
    code = code.replace("\r\n", "\n").replace("\r", "\n")
    code = code.rstrip() + "\n"
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def swiss_dataset() -> pd.DataFrame:
    # Get swiss dataset via statsmodels (R dataset mirror)
    d = sm.datasets.get_rdataset("swiss").data
    return d


def ols_summary_table(model) -> pd.DataFrame:
    params = model.params
    bse = model.bse
    tvals = model.tvalues
    pvals = model.pvalues
    out = pd.DataFrame({
        "coef": params,
        "std_err": bse,
        "t": tvals,
        "p_value": pvals,
    })
    return out


def compute_vif(X: pd.DataFrame) -> pd.DataFrame:
    Xc = sm.add_constant(X, has_constant="add")
    rows = []
    for i, col in enumerate(Xc.columns):
        if col == "const":
            continue
        rows.append({"feature": col, "VIF": float(variance_inflation_factor(Xc.values, i))})
    return pd.DataFrame(rows).sort_values("VIF", ascending=False)


def save_scatter(path: Path, x: pd.Series, y: pd.Series, title: str, xlabel: str, ylabel: str) -> None:
    plt.figure(figsize=(6.2, 4.2))
    sns.scatterplot(x=x, y=y, s=35, alpha=0.75)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(path, dpi=200)
    plt.close()


def task1(doc: Document, assets: Path, counters: Counters) -> None:
    add_center_heading(doc, "Задача 1", level=1)
    doc.add_paragraph(
        "Необходимо загрузить набор данных Swiss и выполнить базовый описательный анализ выбранных переменных, "
        "а также построить простые линейные регрессии вида y = ax + b и оценить их качество."
    )

    df = swiss_dataset()
    y = "Education"
    regressors = ["Infant.Mortality", "Catholic"]

    # Summary stats table
    stats_rows = []
    for col in [y] + regressors:
        s = df[col]
        stats_rows.append({
            "Переменная": col,
            "Среднее": float(s.mean()),
            "Дисперсия": float(s.var(ddof=1)),
            "СКО": float(s.std(ddof=1)),
        })
    stats_df = pd.DataFrame(stats_rows)
    add_table(doc, stats_df, "Описательная статистика (Swiss)", counters)

    # Scatter plots
    p1 = assets / "task1_scatter_edu_infant.png"
    save_scatter(p1, df[regressors[0]], df[y], "Education vs Infant.Mortality", regressors[0], y)
    add_picture(doc, p1, "Диаграмма рассеяния Education от Infant.Mortality", counters)

    p2 = assets / "task1_scatter_edu_catholic.png"
    save_scatter(p2, df[regressors[1]], df[y], "Education vs Catholic", regressors[1], y)
    add_picture(doc, p2, "Диаграмма рассеяния Education от Catholic", counters)

    # Two simple regressions
    for x in regressors:
        X = sm.add_constant(df[[x]])
        m = sm.OLS(df[y], X).fit()
        summ = ols_summary_table(m)
        add_table(
            doc,
            summ,
            f"Параметры модели OLS: {y} ~ {x} (R²={m.rsquared:.3f})",
            counters,
        )

        doc.add_paragraph(
            f"По модели {y} ~ {x}: коэффициент детерминации R²={m.rsquared:.3f}. "
            "При малом R² линейная зависимость выражена слабо; знак коэффициента при регрессоре задаёт "
            "направление связи (положительная/отрицательная)."
        )

    doc.add_paragraph(
        "Вывод: в задаче выполнены описательные статистики и построены две простые линейные регрессии. "
        "Качество оценено по R² и p-value коэффициентов."
    )


def task2(doc: Document, assets: Path, counters: Counters) -> None:
    doc.add_page_break()
    add_center_heading(doc, "Задача 2", level=1)
    doc.add_paragraph(
        "Необходимо проверить отсутствие сильной линейной зависимости между регрессорами, построить "
        "множественную регрессию и подобрать улучшенную модель с трансформациями/взаимодействиями."
    )

    df = swiss_dataset()
    y = "Education"
    regs = ["Catholic", "Agriculture", "Infant.Mortality"]

    # Pairwise R^2 between regressors (as in the R script)
    pairs = [("Catholic", "Agriculture"), ("Catholic", "Infant.Mortality"), ("Agriculture", "Infant.Mortality")]
    rows = []
    for a, b in pairs:
        X = sm.add_constant(df[[b]])
        m = sm.OLS(df[a], X).fit()
        rows.append({"Модель": f"{a} ~ {b}", "R²": float(m.rsquared)})
    add_table(doc, pd.DataFrame(rows), "Проверка линейной зависимости между регрессорами (R²)", counters)

    # Base multiple regression
    X = sm.add_constant(df[regs])
    m0 = sm.OLS(df[y], X).fit()
    add_table(doc, ols_summary_table(m0), f"Множественная регрессия: {y} ~ {', '.join(regs)} (Adj.R²={m0.rsquared_adj:.3f})", counters)

    vif0 = compute_vif(df[regs])
    add_table(doc, vif0, "VIF для базовой модели", counters)

    # A compact "improved" model consistent with the R file idea: log(Agriculture)
    df2 = df.copy()
    df2["log_Agriculture"] = np.log(df2["Agriculture"])
    X1 = sm.add_constant(df2[["log_Agriculture"]])
    m1 = sm.OLS(df2[y], X1).fit()
    add_table(doc, ols_summary_table(m1), f"Упрощённая модель: {y} ~ log(Agriculture) (Adj.R²={m1.rsquared_adj:.3f})", counters)

    # Plot fitted vs actual for the best simple model
    p = assets / "task2_fit_log_agriculture.png"
    plt.figure(figsize=(6.2, 4.2))
    plt.scatter(m1.fittedvalues, df2[y], s=35, alpha=0.75)
    plt.xlabel("Прогноз")
    plt.ylabel("Факт")
    plt.title("Education: факт vs прогноз (log(Agriculture))")
    plt.tight_layout()
    plt.savefig(p, dpi=200)
    plt.close()
    add_picture(doc, p, "Сравнение факта и прогноза для модели Education ~ log(Agriculture)", counters)

    doc.add_paragraph(
        "Вывод: сильной мультиколлинеарности между выбранными регрессорами не обнаружено (R² парных моделей невысокий, VIF умеренный). "
        "Построена базовая множественная регрессия и показан пример улучшенной/упрощённой модели с трансформацией log(Agriculture), "
        "которую удобно интерпретировать и сравнивать по Adj.R²."
    )


def task3(doc: Document, assets: Path, counters: Counters, sav_path: Path) -> None:
    doc.add_page_break()
    add_center_heading(doc, "Задача 3", level=1)
    doc.add_paragraph(
        "В задаче используется SPSS-файл с социально-экономическими признаками. Требуется подготовить данные "
        "(очистка/кодирование/нормализация), построить регрессионные модели и оценить их по R² и диагностическим метрикам (VIF, доверительные интервалы)."
    )

    df, meta = pyreadstat.read_sav(str(sav_path))

    # Select and clean columns similarly to R script
    cols = ["ij13.2", "i_age", "ih5", "i_educ", "status", "ij6.2", "i_marst"]
    data = df[cols].copy()
    data = data.dropna()

    # Salary: numeric conversion + z-score
    salary_raw = pd.to_numeric(data["ij13.2"].astype(str).str.replace(r"[^0-9\-]", "", regex=True), errors="coerce")
    age_raw = pd.to_numeric(data["i_age"].astype(str).str.replace(r"[^0-9\-]", "", regex=True), errors="coerce")
    dur_raw = pd.to_numeric(data["ij6.2"].astype(str).str.replace(r"[^0-9\-]", "", regex=True), errors="coerce")

    data = data.assign(
        salary=(salary_raw - salary_raw.mean()) / salary_raw.std(ddof=1),
        age=(age_raw - age_raw.mean()) / age_raw.std(ddof=1),
        dur=(dur_raw - dur_raw.mean()) / dur_raw.std(ddof=1),
    )

    # sex: ih5 == '1' => 1 else 0
    sex = data["ih5"].astype(str)
    data["sex"] = (sex == "1").astype(int)

    # higher_educ: i_educ in {21,22,23}
    educ = data["i_educ"].astype(str)
    data["higher_educ"] = educ.isin({"21", "22", "23"}).astype(int)

    # status2: status in {1,2}
    status = data["status"].astype(str)
    data["status2"] = status.isin({"1", "2"}).astype(int)

    # marital status dummies (wed1..wed4) based on i_marst
    marst = data["i_marst"].astype(str)
    data["wed1"] = marst.isin({"1", "3"}).astype(int)
    data["wed2"] = (marst == "2").astype(int)
    data["wed3"] = (marst == "4").astype(int)
    data["wed4"] = (marst == "5").astype(int)

    model_df = data[["salary", "age", "sex", "higher_educ", "status2", "dur", "wed1", "wed2", "wed3", "wed4"]].dropna()

    # Descriptive table
    desc = model_df[["salary", "age", "dur"]].agg(["mean", "median", "std"]).T
    add_table(doc, desc.reset_index().rename(columns={"index": "Переменная"}), "Описательные статистики (нормированные признаки)", counters)

    # Base model: salary ~ all
    X = sm.add_constant(model_df.drop(columns=["salary"]))
    m = sm.OLS(model_df["salary"], X).fit()
    add_table(doc, ols_summary_table(m), f"Регрессия: salary ~ признаки (Adj.R²={m.rsquared_adj:.3f})", counters)

    vif = compute_vif(model_df.drop(columns=["salary"]))
    add_table(doc, vif, "VIF для базовой модели", counters)

    # Plot salary distribution
    p = assets / "task3_salary_hist.png"
    plt.figure(figsize=(6.2, 4.2))
    plt.hist(model_df["salary"].values, bins=30)
    plt.title("Распределение нормированной salary")
    plt.xlabel("salary (z-score)")
    plt.ylabel("count")
    plt.tight_layout()
    plt.savefig(p, dpi=200)
    plt.close()
    add_picture(doc, p, "Гистограмма нормированной зарплаты", counters)

    doc.add_paragraph(
        "Вывод: данные очищены от пропусков, выполнено кодирование категориальных признаков и нормализация числовых. "
        "Построена линейная регрессия, оценены коэффициенты, значимость и мультиколлинеарность по VIF."
    )


@dataclass
class NbAsset:
    kind: str  # 'image' | 'table' | 'text'
    title: str
    path: Optional[Path] = None
    df: Optional[pd.DataFrame] = None
    text: Optional[str] = None


def _join_source(src) -> str:
    if isinstance(src, list):
        return "".join(src)
    return str(src or "")


def extract_notebook_assets(nb_path: Path, out_dir: Path, max_tables: int = 10) -> list[NbAsset]:
    nb = json.loads(nb_path.read_text(encoding="utf-8", errors="replace"))
    assets: list[NbAsset] = []
    img_idx = 0
    tbl_idx = 0

    for cell in nb.get("cells", []):
        if cell.get("cell_type") == "markdown":
            md = _join_source(cell.get("source", [])).strip()
            if md:
                # Keep only reasonably short markdown blocks
                assets.append(NbAsset(kind="text", title="Текст", text=md))

        if cell.get("cell_type") != "code":
            continue

        for out in cell.get("outputs", []):
            if not isinstance(out, dict):
                continue
            data = out.get("data", {})
            if not isinstance(data, dict):
                continue

            if "image/png" in data:
                img_data = data["image/png"]
                if isinstance(img_data, list):
                    img_data = "".join(img_data)
                try:
                    raw = base64.b64decode(img_data)
                except Exception:
                    continue
                img_idx += 1
                img_path = out_dir / f"{nb_path.stem}_img_{img_idx}.png"
                img_path.write_bytes(raw)
                assets.append(NbAsset(kind="image", title=f"Иллюстрация из {nb_path.name}", path=img_path))

            if "text/html" in data and tbl_idx < max_tables:
                html = data["text/html"]
                html = _join_source(html)
                if "<table" in html:
                    # Attempt to parse the first table
                    m = re.search(r"(<table[\s\S]*?</table>)", html)
                    table_html = m.group(1) if m else html
                    try:
                        dfs = pd.read_html(table_html)
                    except Exception:
                        continue
                    if dfs:
                        tbl_idx += 1
                        assets.append(NbAsset(kind="table", title=f"Таблица из {nb_path.name}", df=dfs[0]))

    return assets


def task_from_notebook(doc: Document, nb_path: Path, assets_dir: Path, counters: Counters, task_title: str) -> None:
    doc.add_page_break()
    add_center_heading(doc, task_title, level=1)
    extracted = extract_notebook_assets(nb_path, assets_dir)

    # Add a short intro
    doc.add_paragraph(
        f"В этом разделе приведены результаты выполнения {task_title.lower()} на основе ноутбука {nb_path.name}: "
        "краткие пояснения, таблицы и рисунки."
    )

    for item in extracted:
        if item.kind == "text" and item.text:
            # Render markdown simply: strip leading '#', keep paragraphs
            text = re.sub(r"^#+\s*", "", item.text, flags=re.MULTILINE)
            for para in [p.strip() for p in text.split("\n\n") if p.strip()]:
                doc.add_paragraph(para)
            doc.add_paragraph("")
        elif item.kind == "table" and item.df is not None:
            add_table(doc, item.df, item.title, counters, max_rows=15)
        elif item.kind == "image" and item.path is not None:
            add_picture(doc, item.path, item.title, counters)


def add_references(doc: Document) -> None:
    doc.add_page_break()
    add_center_heading(doc, "Список использованной литературы", level=1)
    refs = [
        "[1] scikit-learn: Clustering — документация и пользовательское руководство.",
        "[2] Jolliffe I. T. Principal Component Analysis (PCA): основные идеи и практика.",
        "[3] van der Maaten L., Hinton G. Visualizing Data using t-SNE.",
        "[4] Ester M. et al. A density-based algorithm for discovering clusters (DBSCAN).",
        "[5] Bishop C. M. Pattern Recognition and Machine Learning: Gaussian Mixture Models и EM.",
    ]
    for r in refs:
        doc.add_paragraph(r)


def add_appendices(doc: Document, nir_dir: Path) -> None:
    doc.add_page_break()
    add_center_heading(doc, "Приложения", level=1)
    doc.add_paragraph("Ниже приведены листинги исходного кода (R/Notebook).")

    appendix = [
        nir_dir / "task1.R",
        nir_dir / "task2_edited.R",
        nir_dir / "task3.R",
    ]

    for p in appendix:
        if not p.exists():
            continue
        add_center_heading(doc, f"Код: {p.name}", level=2)
        add_code_block(doc, read_text_file(p))


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate full NIR report with figures/tables.")
    parser.add_argument("--group", default="КМБО-12-24")
    parser.add_argument("--student", default="ФамилияИО")
    parser.add_argument("--variant", default="17")
    parser.add_argument("--semester", default="2семестр")
    parser.add_argument("--out", default=None)
    args = parser.parse_args()

    nir_dir = Path(__file__).resolve().parent
    assets = nir_dir / "report_assets"
    assets.mkdir(exist_ok=True)

    out_name = args.out or f"НИР_{safe_filename(args.group)}_{safe_filename(args.student)}_{safe_filename(args.semester)}_FULL.docx"
    out_path = Path(out_name)
    if not out_path.is_absolute():
        out_path = nir_dir / out_path

    doc = Document()
    set_base_styles(doc)
    counters = Counters()

    # Title page
    add_center_heading(doc, "ОТЧЁТ ПО НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ", level=0)
    p = doc.add_paragraph(f"Вариант: {args.variant}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Группа: {args.group}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Студент: {args.student}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC placeholder
    add_center_heading(doc, "Оглавление", level=1)
    doc.add_paragraph(
        "Оглавление формируется автоматически по заголовкам. В Word: Ссылки → Оглавление → Автособираемое, затем обновите."
    )

    # Tasks
    doc.add_page_break()
    task1(doc, assets, counters)
    task2(doc, assets, counters)

    sav_path = nir_dir / "r13i_os26b.sav"
    task3(doc, assets, counters, sav_path)

    # Tasks 4–6 from notebooks (extraction of embedded outputs)
    task_from_notebook(doc, nir_dir / "task4_variant17_colab.ipynb", assets, counters, "Задача 4")
    task_from_notebook(doc, nir_dir / "task5_ccdata_pca.ipynb", assets, counters, "Задача 5")
    task_from_notebook(doc, nir_dir / "task6_ccdata_clustering.ipynb", assets, counters, "Задача 6")

    # Conclusion
    doc.add_page_break()
    add_center_heading(doc, "Заключение", level=1)
    doc.add_paragraph(
        "В ходе НИР выполнены задачи 1–6: построены и проанализированы регрессионные модели на данных Swiss и SPSS-наборе, "
        "а также проведена предобработка и исследовательский анализ данных держателей кредитных карт с применением PCA, t-SNE и методов кластеризации. "
        "Сравнение алгоритмов кластеризации выполнено по внутренним метрикам и метрикам сходства разбиений (ARI/NMI), "
        "а также оценена доля объектов, которые трудно отнести к кластеру."
    )

    add_references(doc)
    add_appendices(doc, nir_dir)

    doc.save(out_path)
    print(f"OK: saved {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
