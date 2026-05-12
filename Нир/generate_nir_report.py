#!/usr/bin/env python3
"""Generate a NIR report (.docx) following the provided guidelines.

What this script does:
- Creates a Word document with the required structure: title, TOC placeholder, tasks 1–6, conclusion,
  references placeholder (>=5), appendices with code.
- Pulls code automatically from:
  - R scripts (*.R) for tasks 1–3
  - Notebooks (*.ipynb) for tasks 4–6 (code cells only)

Notes:
- Word TOC in python-docx is not a real auto-updating field; we insert a placeholder text.
  In Word you can generate/update TOC from headings.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Iterable


def ensure_package(import_name: str, pip_name: str | None = None) -> None:
    """Ensure an import is available; install via pip if missing."""
    try:
        __import__(import_name)
    except ImportError:
        pkg = pip_name or import_name
        subprocess.check_call([sys.executable, "-m", "pip", "-q", "install", pkg])


ensure_package("docx", "python-docx")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402


def set_base_styles(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 14) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)


def add_center_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_title(doc: Document, lines: Iterable[str]) -> None:
    for i, line in enumerate(lines):
        if not line.strip():
            doc.add_paragraph("")
            continue
        if i == 0:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc: Document, code: str) -> None:
    """Add code as a visually separated block.

    python-docx has limited support for rich code formatting; we keep it simple:
    - monospaced font
    - a 1-cell table to simulate a frame
    """

    # Normalize line endings and trim trailing whitespace lines.
    code = code.replace("\r\n", "\n").replace("\r", "\n")
    code = code.rstrip() + "\n"

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_ipynb_code(path: Path) -> str:
    nb = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    cells = nb.get("cells", [])
    parts: list[str] = []
    code_cell_idx = 0

    for cell in cells:
        if cell.get("cell_type") != "code":
            continue
        src = cell.get("source", [])
        if isinstance(src, list):
            text = "".join(src)
        else:
            text = str(src)
        text = text.strip("\n")
        if not text.strip():
            continue
        code_cell_idx += 1
        parts.append(f"# ---- Notebook cell {code_cell_idx} ----\n{text}\n")

    return "\n".join(parts).strip() + "\n"


def safe_filename(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^0-9A-Za-zА-Яа-я_\-]+", "", s)
    return s


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate NIR report docx (tasks 1–6 + appendices).")
    parser.add_argument("--group", default="КМБО-12-24", help="Group, e.g. КМБО-12-24")
    parser.add_argument("--student", default="ФамилияИО", help="Student short name, e.g. ИвановИИ")
    parser.add_argument("--semester", default="2семестр", help="Semester label")
    parser.add_argument("--variant", default="17", help="Variant number")
    parser.add_argument(
        "--out",
        default=None,
        help="Output docx path. Default: НИР_<group>_<student>_<semester>.docx",
    )

    args = parser.parse_args()

    out_path = Path(args.out) if args.out else Path(f"НИР_{safe_filename(args.group)}_{safe_filename(args.student)}_{safe_filename(args.semester)}.docx")
    if not out_path.is_absolute():
        out_path = Path(__file__).resolve().parent / out_path

    nir_dir = Path(__file__).resolve().parent

    # Expected inputs (can be adjusted later if filenames change)
    task_sources = {
        1: [nir_dir / "task1.R"],
        2: [nir_dir / "task2_edited.R"],
        3: [nir_dir / "task3.R"],
        4: [nir_dir / "task4_variant17_colab.ipynb"],
        5: [nir_dir / "task5_ccdata_pca.ipynb"],
        6: [nir_dir / "task6_ccdata_clustering.ipynb"],
    }

    doc = Document()
    set_base_styles(doc)

    # Title page (minimal, user can replace with their official title page template)
    add_page_title(
        doc,
        [
            "ОТЧЁТ ПО НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ",
            "",
            f"Вариант: {args.variant}",
            f"Группа: {args.group}",
            f"Студент: {args.student}",
            "",
            f"Дата: {datetime.now().strftime('%d.%m.%Y')}",
        ],
    )
    doc.add_page_break()

    # TOC placeholder
    add_center_heading(doc, "Оглавление", level=1)
    doc.add_paragraph(
        "Оглавление формируется автоматически по заголовкам. "
        "В Word: Ссылки → Оглавление → Автособираемое, затем при необходимости обновите."
    )
    doc.add_page_break()

    # Main part: tasks 1–6
    for task_no in range(1, 7):
        add_center_heading(doc, f"Задача {task_no}", level=1)
        doc.add_paragraph("Постановка задачи: (вставьте формулировку из задания/варианта)")
        doc.add_paragraph("Краткая теория/метод: (по желанию, 3–7 предложений)")
        doc.add_paragraph("Решение: (описание шагов, таблицы/рисунки/метрики)")
        doc.add_paragraph("Вывод: (основные результаты по задаче, 2–5 предложений)")

        files = [p for p in task_sources.get(task_no, []) if p.exists()]
        if files:
            doc.add_paragraph("Файлы с кодом/ноутбуки для задачи:")
            for p in files:
                doc.add_paragraph(f"- {p.name}", style="List Bullet")
        else:
            doc.add_paragraph("Файлы с кодом для задачи: (не найдены автоматически — проверьте имена файлов)")

        doc.add_page_break()

    # Conclusion
    add_center_heading(doc, "Заключение", level=1)
    doc.add_paragraph(
        "Краткая характеристика выводов по каждой задаче (можно 1–2 предложения на задачу)."
    )
    doc.add_page_break()

    # References (placeholder, needs >=5)
    add_center_heading(doc, "Список использованной литературы", level=1)
    doc.add_paragraph(
        "Добавьте не менее 5 источников, оформленных по ГОСТ. Нумерация — по мере появления в тексте."
    )
    for i in range(1, 6):
        doc.add_paragraph(f"[{i}] ", style="List Number")
    doc.add_page_break()

    # Appendices: code listings
    add_center_heading(doc, "Приложения", level=1)
    doc.add_paragraph(
        "Ниже приведены листинги кода для решения задач (код вставлен текстом, как требуется в методичке)."
    )

    appendix_idx = 0
    for task_no in range(1, 7):
        files = [p for p in task_sources.get(task_no, []) if p.exists()]
        if not files:
            continue

        for src_path in files:
            appendix_idx += 1
            add_center_heading(doc, f"Приложение {appendix_idx}. Код к задаче {task_no}: {src_path.name}", level=2)

            try:
                if src_path.suffix.lower() == ".ipynb":
                    code = extract_ipynb_code(src_path)
                else:
                    code = read_text_file(src_path)
            except Exception as e:
                doc.add_paragraph(f"Не удалось извлечь код из {src_path.name}: {e}")
                continue

            if not code.strip():
                doc.add_paragraph("(Код пустой)")
                continue

            add_code_block(doc, code)

    doc.save(out_path)
    print(f"OK: report saved to {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
